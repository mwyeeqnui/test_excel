''' 作者：崔鹏飞
    版本号：不管是几吧！
    功能：我也不太懂！
'''

import pprint
from tkinter import *
from tkinter import filedialog
import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Length, Pt, RGBColor


def handle_excel_gui():
    root = Tk()
    root.title("处理excel")
    root.geometry('425x280')
    excel_path = StringVar()
    doc_path = StringVar()

    def get_excel():
        excel_name = filedialog.askopenfilename(filetypes=[("XLSX", '.xlsx')])
        excel_path.set(excel_name)

    def tar_dir():
        doc_name = filedialog.askdirectory()
        doc_path.set(doc_name)

    def handle_excel(dir_path):

        wb = openpyxl.load_workbook(dir_path)
        get_sheet_name = wb.get_sheet_names()[0]  # 只能取出EXCEL中的第一个工作表
        sheet = wb.get_sheet_by_name(get_sheet_name)
        r_value = []  # 用于放入R列值的list，减少循环次数--> 临时list
        r_start_int = 2  # range中的最小值，方法循环
        end_list = []  # 最终要装入所有数据的list

        #  对EXCEL排序，取所有需要的值到列表。
        for row in range(r_start_int, sheet.max_row + 1):  # 循环EXCEL的最大行数次。
            if sheet['D{0}'.format(
                    row)].value in r_value:  # 如果D0的值在list中，跳入下次循环
                continue
            else:  # 否则加入end_list中
                r_value.append(sheet['D{0}'.format(row)].value)
            for r_row in range(r_start_int + 1,
                               sheet.max_row + 1):  # 循环EXCEL的最大行数次减1，跳过比对过的行。
                if sheet['D{0}'.format(r_row)].value == sheet['D{0}'.format(
                        row)].value:  #  如果D列的值相同，将以下的值插入end_list
                    end_list.append([
                        sheet['D{0}'.format(r_row)].value,
                        sheet['H{0}'.format(r_row)].value,
                        sheet['I{0}'.format(r_row)].value,
                        sheet['J{0}'.format(r_row)].value,
                        sheet['M{0}'.format(r_row)].value,
                        sheet['N{0}'.format(r_row)].value,
                        sheet['O{0}'.format(r_row)].value,
                        sheet['P{0}'.format(r_row)].value,
                        sheet['R{0}'.format(r_row)].value
                    ])

        #  按照list的第0个元素即D列对列表中的值切片，分段
        d_test_list = []
        d_slice_list = []  # 存放对D列切片后的list
        d_int_test = 0
        for i in end_list:
            if i[0] not in d_test_list:
                d_test_list.append(i[0])
                d_slice_list.append(end_list[d_int_test:end_list.index(i)])
                d_int_test = end_list.index(i)
            else:
                continue
        else:
            d_slice_list.append(end_list[d_int_test:len(end_list)])

        #  list第0个元素切完之后，按照list第7个元素对list再次进行切片。
        p_end_list = []
        for every_list_int in range(len(d_slice_list)):
            #  对所有的d_slice_list的第7个元素进行排序
            p_int = 0
            every_sorted = sorted(
                d_slice_list[every_list_int], key=lambda x: x[8], reverse=True)
            p_check_list = []
            for (offset, item) in enumerate(every_sorted):  #取出每个list和下标
                # print(item)
                if item[8] not in p_check_list:  #如果下标7不存在于list中，就添加
                    p_check_list.append(item[8])
                    p_end_list.append(every_sorted[p_int:offset])
                    p_int = offset
                else:  # 如果存在就跳入下一个循环
                    continue
            else:
                p_end_list.append(every_sorted[p_int:len(every_sorted)])

        for (offset, i) in enumerate(p_end_list):
            if i:
                pass
            else:
                p_end_list.remove(i)

        return p_end_list

    def to_docx():
        excel_list = handle_excel(en_excel.get())
        item_int = 0
        test_int = 0
        for (offset, item) in enumerate(excel_list):
            if item and len(item) <= 40:
                type_int = str(len(item))
                str_pact = item[0][5]
                str_type = item[0][0]

                document = Document()
                document.styles['Normal'].font.name = u'宋体'
                document.styles['Normal'].font.size = Pt(14)
                document.styles['Normal']._element.rPr.rFonts.set(
                    qn('w:eastAsia'), u'宋体')

                p = document.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph_format = p.paragraph_format
                paragraph_format.line_spacing = 1
                run_p = p.add_run(u'湖南宏达工程造价咨询有限公司')
                run_p.font.color.rgb = RGBColor(255, 0, 0)
                run_p.font.name = u'楷体'
                run_p.font.size = Pt(28)
                run_p.bold = True

                p0 = document.add_paragraph()
                p0.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph_format = p0.paragraph_format
                paragraph_format.line_spacing = Pt(18)
                run_p0 = p0.add_run(u'湘宏达（包移）审字[2017]CB-006号')
                run_p0.font.name = u'楷体'
                run_p0.font.size = Pt(16)
                run_p0.bold = True

                document.add_paragraph()  #  空行

                p1 = document.add_paragraph()
                p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                run1 = p1.add_run(u'关于中国移动通信集团内蒙古有限公司包头业务区')
                run1.font.size = Pt(15)
                run1.bold = True

                # 对相关字段设置红色
                run_type = p1.add_run(str_type)
                run_type.font.size = Pt(15)
                run_type.bold = True

                run2 = p1.add_run(str_pact + '等')
                run2.font.size = Pt(15)
                run2.bold = True

                # 对相关字段设置红色
                run_int = p1.add_run(type_int)
                run_int.font.size = Pt(15)
                run_int.bold = True

                run3 = p1.add_run(u'项工程结算审核报告')
                run3.font.size = Pt(15)
                run3.bold = True

                document.add_paragraph()  #  空行

                p2 = document.add_paragraph()
                run_p2 = p2.add_run('中国移动通信集团内蒙古有限公司：')
                paragraph_format = p2.paragraph_format
                paragraph_format.line_spacing = 1.25
                run_p2.font.size = Pt(16)
                run_p2.bold = True

                p3 = document.add_paragraph()
                paragraph_format = p3.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                run_p3 = p3.add_run('受中国移动通信集团内蒙古有限公司的委托，湖南宏达工程造价咨询有限公司对包头业务区')
                run_p3.font.size = Pt(14)
                run_p3_1 = p3.add_run(str_type)
                run_p3_1.font.size = Pt(14)
                run_p3_2 = p3.add_run(str_pact + '等')
                run_p3_2.font.size = Pt(14)
                run_p3_3 = p3.add_run(type_int)
                run_p3_3.font.size = Pt(14)
                run_p3_4 = p3.add_run('项工程结算进行了审核。')
                run_p3_4.font.size = Pt(14)

                p4 = document.add_paragraph('')
                paragraph_format = p4.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                run_p4 = p4.add_run(
                    '贵公司的责任是对提供资料的真实性、合法性、完整性负责。我公司的责任是在贵公司提供结算资料的基础上进行结算审核并发表审核意见。'
                )
                run_p4.font.size = Pt(14)

                p5 = document.add_paragraph('')
                paragraph_format = p5.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                run_p4 = p5.add_run(
                    '在审核过程中，我们结合上述项目的实际情况，实施了包括现场察勘、询问、审核结算资料、复核、和有关工程技术人员交换意见等，我们认为必要的审核程序。在贵公司的大力支持和相关人员的密切配合下，审核工作已结束。现将审核情况报告如下：'
                )
                run_p4.font.size = Pt(14)

                document.add_paragraph()  # 空行

                document.add_heading('一 、工程概况', level=2)

                zong_jing_e = 0.00
                song_shen_e = 0.00
                shen_ding_e = 0.00
                for every_list in item:
                    zong_jing_e += float(every_list[1])
                    song_shen_e += float(every_list[2])
                    shen_ding_e += float(every_list[3])
                sheng_jian = song_shen_e - shen_ding_e
                avg_shen_jian = (sheng_jian / song_shen_e) * 100
                p6 = document.add_paragraph()
                paragraph_format = p6.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                run_p6 = p6.add_run(
                    '中国移动通信集团内蒙古有限公司包头业务区' + str_type + str_pact + '等' +
                    type_int +
                    '项工程合同金额共{:.2f}元；报审金额为{:.2f}元；审定金额为{:.2f}元；审减金额{:.2f}元，平均审减率为{:.2f}%.本报告将各施工单位所做工程情况简介如下：'
                    .format(zong_jing_e, song_shen_e, shen_ding_e, sheng_jian,
                            avg_shen_jian))
                run_p6.font.size = Pt(14)

                #  取值
                sort_company = sorted(item, key=lambda x: x[7])
                company_int = 0
                company_list = []
                con_int = 0
                sec_songshen_e = 0.00
                sec_shending_e = 0.00
                for (_offset, list_company) in enumerate(sort_company):

                    if list_company[7] not in company_list:
                        # 如果不在列表中就添加进去
                        con_int = 1
                        company_int += 1
                        company_list.append(list_company[7])
                        p7 = document.add_paragraph()
                        run_p7 = p7.add_run(
                            str(1) + '.{0}'.format(company_int) +
                            list_company[7])
                        run_p7.bold = True
                        run_p7.font.size = Pt(14)

                        p8 = document.add_paragraph()
                        run_p8 = p8.add_run(
                            '1.{0}.1工程名称：{1}）{2}（合同编号：{3}）。'.format(
                                company_int, con_int, list_company[5],
                                list_company[4]))
                        paragraph_format = p8.paragraph_format
                        paragraph_format.first_line_indent = Inches(0.5)
                        run_p8.font.size = Pt(14)

                        # 循环取值运算工程送审、审定金额、审减等。
                        p9 = document.add_paragraph()
                        paragraph_format = p9.paragraph_format
                        paragraph_format.first_line_indent = Inches(0.5)
                        for test_item in sort_company:
                            if list_company[7] == test_item[7]:
                                sec_songshen_e += float(test_item[2])
                                sec_shending_e += float(test_item[3])
                        sec_shenjian_e = sec_songshen_e - sec_shending_e
                        run_p9 = p9.add_run(
                            '1.{}.2 工程造价：工程送审{:.2f}元，审定金额为{:.2f}元，审减造价{:.2f}元。'
                            .format(company_int, sec_songshen_e,
                                    sec_shending_e, sec_shenjian_e))

                        sec_songshen_e = 0.00
                        sec_shending_e = 0.00
                        sec_shenjian_e = 0.00

                    else:
                        con_int += 1
                        run_p8 = p8.add_run('{0}）{1}（合同编号：{2}）。'.format(
                            con_int, list_company[5], list_company[4]))
                        paragraph_format = p8.paragraph_format
                        paragraph_format.first_line_indent = Inches(0.5)

                p10 = document.add_paragraph()
                run_p10 = p10.add_run('1.{} 工程建设批复情况：'.format(company_int + 1))
                run_p10.bold = True

                p11 = document.add_paragraph()
                paragraph_format = p11.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                run_p11 = p11.add_run('1、批复情况：')
                run_p11.bold = True
                run_p11_2 = p11.add_run('本报告涉及立项批复内容较多，且均为会议纪要，在报告中不予详书。')

                p12 = document.add_paragraph()
                paragraph_format = p12.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                run_p12 = p12.add_run('2、工程建设单位：')
                run_p12.bold = True
                run_p12_2 = p12.add_run('中国移动通信集团内蒙古有限公司包头分公司')

                p13 = document.add_paragraph()
                paragraph_format = p13.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                run_p13 = p13.add_run('3、工程建设单位：')
                run_p13.bold = True
                run_p13_2 = p13.add_run('无')

                document.add_heading('二 、审计内容', level=2)

                p14 = document.add_paragraph()
                paragraph_format = p14.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p14.add_run('1、结算方式是否符合合同约定。')

                p15 = document.add_paragraph()
                paragraph_format = p15.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p15.add_run('2、结算工程量与实际工程量的误差情况。')

                p16 = document.add_paragraph()
                paragraph_format = p16.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p16.add_run('3、工程量是否符合规定的计算规则，数量是否准确。')

                p17 = document.add_paragraph()
                paragraph_format = p17.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p17.add_run('4、结算单价的套用是否合理。')

                p18 = document.add_paragraph()
                paragraph_format = p18.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p18.add_run('5、工程取费是否执行相应的计算基数和费率标准等。')

                document.add_heading('三、审计依据', level=2)

                p19 = document.add_paragraph()
                paragraph_format = p19.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p19.add_run(
                    '1、中国移动通信集团内蒙古有限公司与湖南宏达工程造价咨询有限公司签订的《中国移动通信集团公司内蒙古有限公司委托建设项目结算审计框架合同》合同'
                )

                p20 = document.add_paragraph()
                paragraph_format = p20.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p20.add_run('2、国家、行业相关文件及规定。')

                p21 = document.add_paragraph()
                paragraph_format = p21.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p21.add_run('3、建设单位提供的资料：')

                p22 = document.add_paragraph()
                paragraph_format = p22.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p22.add_run(' （1）原报工程结算书；')

                p23 = document.add_paragraph()
                paragraph_format = p23.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p23.add_run(' （2）施工图；')

                p24 = document.add_paragraph()
                paragraph_format = p24.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p24.add_run(' （3）施工合同；')

                p25 = document.add_paragraph()
                paragraph_format = p25.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p25.add_run(' （4）竣工技术文件；')

                p26 = document.add_paragraph()
                paragraph_format = p26.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p26.add_run(' （5）工程签证资料；')

                p27 = document.add_paragraph()
                paragraph_format = p27.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p27.add_run(' （6）工程定价资料；')

                p28 = document.add_paragraph()
                paragraph_format = p28.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p28.add_run(' （7）投标书、招标文件、投标文件。')

                p29 = document.add_paragraph()
                paragraph_format = p29.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p29.add_run('4、审计定额依据：')

                p30 = document.add_paragraph()
                paragraph_format = p30.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p30.add_run(
                    '''《通信建设工程概预算编制配套法规文件汇编（2008.5）》、《通信建设工程价款结算办法》、《通信建设工程预算定额（全册）》、内蒙古包头市材料价格信息及移动公司关于工程结算的有关规定。'''
                )

                p31 = document.add_paragraph()
                paragraph_format = p31.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p31.add_run('5、现场勘察记录。')

                document.add_heading('四、双方责任', level=2)

                p32 = document.add_paragraph()
                run_32 = p32.add_run('4.1 建设单位责任')
                run_32.bold = True

                p33 = document.add_paragraph()
                paragraph_format = p33.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p33.add_run('1、建设单位通过审计委托书向审计机构委托工程建设项目审计，并向审计机构提供工程审计的全部资料。')

                p34 = document.add_paragraph()
                paragraph_format = p34.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p34.add_run(
                    '2、建设单位应对开展审核工作给予充分的合作，提供必要的工作条件，并按审计机构要求，及时提供与委托内容相关的完整的建设项目有关文件、规定、合同、图纸、预结算等资料，并对其真实性、合法性和完整性负责。'
                )

                p35 = document.add_paragraph()
                paragraph_format = p35.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p35.add_run(
                    '3、审核过程中需进一步核对原始资料或计算依据的，建设单位及其施工单位应提供必要的工作条件及合作，凡需施工单位签认的基础资料由建设单位协助办理。'
                )

                p33 = document.add_paragraph()
                paragraph_format = p33.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p33.add_run(
                    '4、工程结算审核结果，凡需建设单位、审计机构双方和施工单位合议的，由建设单位组织协调，施工单位可以充分发表意见，对审核确认的事项，施工单位无理拒签、又不参加合议的，不影响审计机构出具审核报告，由此产生的后果由责任者承担。'
                )

                p34 = document.add_paragraph()
                paragraph_format = p34.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p34.add_run(
                    '5、由于建设单位违反约定事项，所提供的审核资料不真实，不合法、不完整、不及时，致使审核结果不当或延期提交审核报告，由建设单位承担责任。'
                )

                p35 = document.add_paragraph()
                paragraph_format = p35.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p35.add_run('6、建设单位应正确使用审计报告，由于使用不当所造成的后果，与审计机构无关。')

                p36 = document.add_paragraph()
                paragraph_format = p36.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p36.add_run('7、建设单位应按照约定的条件，及时足额支付审计费。')

                p37 = document.add_paragraph()
                run_37 = p37.add_run('4.2审计机构责任')
                run_37.bold = True

                p38 = document.add_paragraph()
                paragraph_format = p38.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p38.add_run('1、工程结算审核范围')

                p39 = document.add_paragraph()
                paragraph_format = p39.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p39.add_run('2、施工方工程施工费、安装费等的审核。')

                p40 = document.add_paragraph()
                paragraph_format = p40.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p40.add_run(
                    '3、审计机构应按照法规和准则的要求，根据建设单位提供的各项建设项目文件和工程资料，严格实施主要的审核程序，发表审核意见，按约定的时间出具真实合法的审核报告。'
                )

                p41 = document.add_paragraph()
                paragraph_format = p41.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p41.add_run(
                    '4、检查弊端不属于一般审计工作范围，但在审计过程如发现被审计单位在基建工程管理，财务管理和财产物质管理方面存在问题，导致有产生重大弊端的可能，审计机构应将其情况报告建设单位。在审计过程中，如发现被审计单位的内部控制有重大缺陷，应将情况报告建设单位。'
                )

                p42 = document.add_paragraph()
                paragraph_format = p42.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p42.add_run(
                    '5、审计机构应按照约定的审计时限、安排足够的审计人员到达审计现场开展审计工作，如不能及时安排审计人员进行审计，则建设单位有权安排其他审计单位进行审计。'
                )

                p43 = document.add_paragraph()
                paragraph_format = p43.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p43.add_run(
                    '6、现场审计：审计机构应对扩容设备安装、综合楼生产楼等土建（包括水、电、暖等）、机房装修及城域网扩容管道、零星工程等工程项目， 100%现场实测、对同一家施工单位施工的光缆线路等单价低、分布广的工程，现场勘查率必须达到20%以上；审计过程中如出现工程量方面的纠纷或施工单位提出要求，必须现场抽查。'
                )

                p44 = document.add_paragraph()
                paragraph_format = p44.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p44.add_run(
                    '7、审计机构应按照中标要求出具审计报告，由于审核人员是采取事后审计方法，并受被审计单位内部控制固有的局限和其他客观因素的制约，以及审核人员在审核中可能未予发现的疏漏。因此，审计机构的审核难免存在某些重要的方面反映失实，但这不能替代、减轻或免除建设单位对工程结算的编制责任。'
                )

                p45 = document.add_paragraph()
                paragraph_format = p45.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p45.add_run('8、审计机构应于每周一报送审计周报，对建设单位委托的审计项目的审计进度、存在问题等情况进行反馈。')

                p46 = document.add_paragraph()
                paragraph_format = p46.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p46.add_run(
                    '9、由于审计机构违反工程造价审核规定和程序，致使审核结果严重失真，审计机构应重新进行审核，由此产生的费用由审计机构自行负责，若给建设单位造成了经济损失，审计机构还应全额赔偿建设单位的损失。'
                )

                document.add_heading('五、审核过程', level=2)

                p47 = document.add_paragraph()
                paragraph_format = p47.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p47.add_run(
                    '我们受贵公司委托，于2017年3月1日至2018年1月15日，对中国移动通信集团内蒙古有限公司包头业务区{0}{1}等{2}项工程进行了审计。在审计过程中，我们采用全面审核的方法，结合工程的实际情况，实施了包括逐项审查结算资料、现场勘察和有关工程技术人员交换意见等我们认为必要的审核程序。'
                    .format(str_type, str_pact, type_int))

                p48 = document.add_paragraph()
                paragraph_format = p48.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p48.add_run('我们根据本项目的特点，制定了审计工作计划。按照审计计划，我们将审计工作分为三个阶段：')

                p49 = document.add_paragraph()
                paragraph_format = p49.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p49.add_run('第一阶段，熟悉工程情况、搜集审计资料工作，并在此期间进行了结算初步审计工作，形成审计明细初稿。')

                p50 = document.add_paragraph()
                paragraph_format = p50.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p50.add_run('搜集的主要资料有：')

                p51 = document.add_paragraph()
                paragraph_format = p51.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p51.add_run('①	工程批复文件、招投标文件；')

                p52 = document.add_paragraph()
                paragraph_format = p52.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p52.add_run('②	工程报审结算书；')

                p53 = document.add_paragraph()
                paragraph_format = p53.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p53.add_run('③	工程施工合同、设计文件；')

                p54 = document.add_paragraph()
                paragraph_format = p54.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p54.add_run('④	工程竣工技术资料（开工报告、完工报告、初验证书、各种涉造价部分文件等）；')

                p55 = document.add_paragraph()
                paragraph_format = p55.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p55.add_run(
                    '第二阶段，实施了100%现场勘察、结算核对工作，对现场安装工程量实行点单型核查，并与施工单位、监理单位、建设单位交换了意见，依据施工合同对结算方式进行了确认，达成共识后，确认工程结算审核定案表。'
                )

                p56 = document.add_paragraph()
                paragraph_format = p56.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p56.add_run('第三阶段，形成审计资料终稿，拟订审计报告初稿，与各方交换意见后，出具正式审计报告。')

                document.add_heading('六、审核结果', level=2)

                p57 = document.add_paragraph()
                run_57 = p57.add_run('6.1 经济性评价')
                run_57.bold = True

                p58 = document.add_paragraph()
                paragraph_format = p58.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p58.add_run(
                    '6.1.1 本报告所含{0}项中国移动通信集团内蒙古有限公司包头业务区{1}等工程的具体情况如下（详见附件一：工程结算费用审核汇总表）：'
                    .format(type_int, str_type))

                p59 = document.add_paragraph()
                paragraph_format = p59.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p59.add_run(
                    '6.1.2 本报告{0}项工程分由{1}家施工单位实施，合同金额共为{2:.2f}元，总送审金额为{3:.2f}元，审计总造价为{4:.2f}元，审减造价{5:.2f}元，平均审减率为{6:.2f}%，审计费用共计{7}元。'
                    .format(type_int, len(company_list), zong_jing_e,
                            song_shen_e, shen_ding_e, sheng_jian,
                            avg_shen_jian, 10403.03))  # 这里有他妈大问题

                p60 = document.add_paragraph()
                paragraph_format = p60.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                run_60 = p60.add_run('6.1.3  这里有他妈大问题！')
                run_60.font.size = Pt(40)

                p61 = document.add_paragraph()
                paragraph_format = p61.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                run_61 = p61.add_run('6.1.4  这里有他妈大问题！')

                p62 = document.add_paragraph()
                run_62 = p62.add_run('6.2 主要核减原因')
                run_62.bold = True

                p63 = document.add_paragraph()
                paragraph_format = p63.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p63.add_run('1、送审工作量与现场工作量不符')

                p64 = document.add_paragraph()
                paragraph_format = p64.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p64.add_run(
                    '例如：商丘市国基建筑安装有限公司送审的青山赛音道渠道装修施工合同工程中，施工单位送审门头贴膜8.766m2，现场实际查勘是4.35m2，审计时都按现场测量情况予以审减。'
                )

                p65 = document.add_paragraph()
                paragraph_format = p65.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                run_65 = p65.add_run('2、定额取值套用错误')

                p66 = document.add_paragraph()
                paragraph_format = p66.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p66.add_run(
                    '例如：山西胜唐通信建设有限公司送审的第二批社会渠道传输改造项目林荫路科技楼他建专营店等2处施工合同中，施工单位送审布放墙壁光缆套取定额子目TXL-049（架设吊线式墙壁光缆），实际现场查看施工单位布放光缆方式是利旧式墙壁吊线，应套取定额子目TXL3-183(城区架设架空光缆)。'
                )

                document.add_heading('七、审核问题', level=2)

                p67 = document.add_paragraph()
                paragraph_format = p67.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p67.add_run('本次审计共发现问题1个，提出1个管理建议。')

                p68 = document.add_paragraph()
                paragraph_format = p68.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p68.add_run(
                    '本报告所涉部分项目为成本类工程，全部采用定额计价法结算，施工单位如果没有专业预算人员，容易发生定额套用错误或者借用定额不符等争议，导致验收、竣工结算、审计时产生不便，既影响施工单位合法效益，又加大建设单位成本管理难度。'
                )

                p69 = document.add_paragraph()
                paragraph_format = p69.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p69.add_run(
                    '管理建议：建议建设单位针对工艺不复杂，材料采购简单且价格容易调查的项目在年度招标时实施清单报价法计价，既方便双方清楚造价的构成又避免后期的争议。建设单位和监理只要负责施工单位施工使用的材料是否与招标要求相符，完成质量、工程量可以在验收是核定确认，方便、明了。'
                )

                document.add_heading('八、其他说明', level=2)

                p70 = document.add_paragraph()
                paragraph_format = p70.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p70.add_run(
                    '1.本报告中所涉建设单位提供的相关依据资料（批复立项资料、施工合同、开工报告、验收报告等）由于在电子档案中已存在，报告不予重复。'
                )

                p71 = document.add_paragraph()
                paragraph_format = p71.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p71.add_run(
                    '2.本报告一式四份，中国移动通信集团内蒙古有限公司及包头分公司共持三份，湖南宏达工程造价咨询有限公司持一份。')

                document.add_paragraph()
                document.add_paragraph()
                document.add_paragraph()
                document.add_paragraph()
                document.add_paragraph()
                document.add_paragraph()
                document.add_paragraph()
                document.add_paragraph()

                p72 = document.add_paragraph()
                run_72 = p72.add_run('附件:')
                run_72.bold = True

                p73 = document.add_paragraph()
                paragraph_format = p73.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p73.add_run('1、附件一：工程结算费用审核汇总表')

                p74 = document.add_paragraph()
                paragraph_format = p74.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p74.add_run('2、附件二：单项工程审计定案表')

                p75 = document.add_paragraph()
                paragraph_format = p75.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p75.add_run('3、附件三：单项工程审定明细表')

                p76 = document.add_paragraph()
                paragraph_format = p76.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p76.add_run('4、附件四：结算审计现场审核记录表')

                p77 = document.add_paragraph()
                paragraph_format = p77.paragraph_format
                paragraph_format.first_line_indent = Inches(0.5)
                p77.add_run('5、附件五：部分现场审计痕迹记录')

                document.add_paragraph()
                document.add_paragraph()
                document.add_paragraph()

                p78 = document.add_paragraph()
                p78.add_run('湖南宏达工程造价咨询有限公司           项目负责人：付红卫')

                p79 = document.add_paragraph()
                p79.add_run(
                    '                                  编审： 彭创  李庆辉  梁湘')

                p80 = document.add_paragraph()
                p80.add_run(
                    '                                  复审： 陈晔  洪艳  胡红梅')

                p81 = document.add_paragraph()
                p81.add_run('地址：长沙市天心区湘府中路369号       审计公司执业印章')

                p82 = document.add_paragraph()
                p82.add_run('     星城荣域综合楼房1310')

                p83 = document.add_paragraph()
                p83.add_run('                                2018年1月25日')

                document.save('{0}\{1}{2}.docx'.format(en_doc.get(), str_type,
                                                       offset))
                list_box.insert(0, '{0}:{1}-->生成完成！'.format(str_type, offset))

            elif len(item) >= 40:
                long_cut_list = [(offset + 1)
                                 for (offset, _item) in enumerate(item)
                                 if (offset + 1) % 40 == 0]
                long_cut_list.append(len(item))
                cut_last_list = []
                last_int = 0
                name_int = 0
                for last_item_int in long_cut_list:
                    cut_last_list.append(item[last_int:last_item_int])
                    last_int = last_item_int

                for (_offset, last_item) in enumerate(cut_last_list):

                    type_int = str(len(last_item))
                    str_pact = item[0][5]
                    str_type = item[0][0]

                    document = Document()
                    document.styles['Normal'].font.name = u'宋体'
                    document.styles['Normal'].font.size = Pt(14)
                    document.styles['Normal']._element.rPr.rFonts.set(
                        qn('w:eastAsia'), u'宋体')

                    p = document.add_paragraph()
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph_format = p.paragraph_format
                    paragraph_format.line_spacing = 1
                    run_p = p.add_run(u'湖南宏达工程造价咨询有限公司')
                    run_p.font.color.rgb = RGBColor(255, 0, 0)
                    run_p.font.name = u'楷体'
                    run_p.font.size = Pt(28)
                    run_p.bold = True

                    p0 = document.add_paragraph()
                    p0.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph_format = p0.paragraph_format
                    paragraph_format.line_spacing = Pt(18)
                    run_p0 = p0.add_run(u'湘宏达（包移）审字[2017]CB-006号')
                    run_p0.font.name = u'楷体'
                    run_p0.font.size = Pt(16)
                    run_p0.bold = True

                    document.add_paragraph()  #  空行

                    p1 = document.add_paragraph()
                    p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    run1 = p1.add_run(u'关于中国移动通信集团内蒙古有限公司包头业务区')
                    run1.font.size = Pt(15)
                    run1.bold = True

                    # 对相关字段设置红色
                    run_type = p1.add_run(str_type)
                    run_type.font.size = Pt(15)
                    run_type.bold = True

                    run2 = p1.add_run(str_pact + '等')
                    run2.font.size = Pt(15)
                    run2.bold = True

                    # 对相关字段设置红色
                    run_int = p1.add_run(type_int)
                    run_int.font.size = Pt(15)
                    run_int.bold = True

                    run3 = p1.add_run(u'项工程结算审核报告')
                    run3.font.size = Pt(15)
                    run3.bold = True

                    document.add_paragraph()  #  空行

                    p2 = document.add_paragraph()
                    run_p2 = p2.add_run('中国移动通信集团内蒙古有限公司：')
                    paragraph_format = p2.paragraph_format
                    paragraph_format.line_spacing = 1.25
                    run_p2.font.size = Pt(16)
                    run_p2.bold = True

                    p3 = document.add_paragraph()
                    paragraph_format = p3.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    run_p3 = p3.add_run(
                        '受中国移动通信集团内蒙古有限公司的委托，湖南宏达工程造价咨询有限公司对包头业务区')
                    run_p3.font.size = Pt(14)
                    run_p3_1 = p3.add_run(str_type)
                    run_p3_1.font.size = Pt(14)
                    run_p3_2 = p3.add_run(str_pact + '等')
                    run_p3_2.font.size = Pt(14)
                    run_p3_3 = p3.add_run(type_int)
                    run_p3_3.font.size = Pt(14)
                    run_p3_4 = p3.add_run('项工程结算进行了审核。')
                    run_p3_4.font.size = Pt(14)

                    p4 = document.add_paragraph('')
                    paragraph_format = p4.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    run_p4 = p4.add_run(
                        '贵公司的责任是对提供资料的真实性、合法性、完整性负责。我公司的责任是在贵公司提供结算资料的基础上进行结算审核并发表审核意见。'
                    )
                    run_p4.font.size = Pt(14)

                    p5 = document.add_paragraph('')
                    paragraph_format = p5.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    run_p4 = p5.add_run(
                        '在审核过程中，我们结合上述项目的实际情况，实施了包括现场察勘、询问、审核结算资料、复核、和有关工程技术人员交换意见等，我们认为必要的审核程序。在贵公司的大力支持和相关人员的密切配合下，审核工作已结束。现将审核情况报告如下：'
                    )
                    run_p4.font.size = Pt(14)

                    document.add_paragraph()  # 空行

                    document.add_heading('一 、工程概况', level=2)

                    zong_jing_e = 0.00
                    song_shen_e = 0.00
                    shen_ding_e = 0.00
                    for every_list in last_item:
                        zong_jing_e += float(every_list[1])
                        song_shen_e += float(every_list[2])
                        shen_ding_e += float(every_list[3])
                    sheng_jian = song_shen_e - shen_ding_e
                    avg_shen_jian = (sheng_jian / song_shen_e) * 100
                    p6 = document.add_paragraph()
                    paragraph_format = p6.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    run_p6 = p6.add_run(
                        '中国移动通信集团内蒙古有限公司包头业务区' + str_type + str_pact + '等' +
                        type_int +
                        '项工程合同金额共{:.2f}元；报审金额为{:.2f}元；审定金额为{:.2f}元；审减金额{:.2f}元，平均审减率为{:.2f}%.本报告将各施工单位所做工程情况简介如下：'
                        .format(zong_jing_e, song_shen_e, shen_ding_e,
                                sheng_jian, avg_shen_jian))
                    run_p6.font.size = Pt(14)

                    #  取值
                    sort_company = sorted(last_item, key=lambda x: x[7])
                    company_int = 0
                    company_list = []
                    con_int = 0
                    sec_songshen_e = 0.00
                    sec_shending_e = 0.00
                    for (_offset, list_company) in enumerate(sort_company):

                        if list_company[7] not in company_list:
                            # 如果不在列表中就添加进去
                            con_int = 1
                            company_int += 1
                            company_list.append(list_company[7])
                            p7 = document.add_paragraph()
                            run_p7 = p7.add_run(
                                str(1) + '.{0}'.format(company_int) +
                                list_company[7])
                            run_p7.bold = True
                            run_p7.font.size = Pt(14)

                            p8 = document.add_paragraph()
                            run_p8 = p8.add_run(
                                '1.{0}.1工程名称：{1}）{2}（合同编号：{3}）。'.format(
                                    company_int, con_int, list_company[5],
                                    list_company[4]))
                            paragraph_format = p8.paragraph_format
                            paragraph_format.first_line_indent = Inches(0.5)
                            run_p8.font.size = Pt(14)

                            # 循环取值运算工程送审、审定金额、审减等。
                            p9 = document.add_paragraph()
                            paragraph_format = p9.paragraph_format
                            paragraph_format.first_line_indent = Inches(0.5)
                            for test_item in sort_company:
                                if list_company[7] == test_item[7]:
                                    sec_songshen_e += float(test_item[2])
                                    sec_shending_e += float(test_item[3])
                            sec_shenjian_e = sec_songshen_e - sec_shending_e
                            run_p9 = p9.add_run(
                                '1.{}.2 工程造价：工程送审{:.2f}元，审定金额为{:.2f}元，审减造价{:.2f}元。'
                                .format(company_int, sec_songshen_e,
                                        sec_shending_e, sec_shenjian_e))

                            sec_songshen_e = 0.00
                            sec_shending_e = 0.00
                            sec_shenjian_e = 0.00

                        else:
                            con_int += 1
                            run_p8 = p8.add_run('{0}）{1}（合同编号：{2}）。'.format(
                                con_int, list_company[5], list_company[4]))
                            paragraph_format = p8.paragraph_format
                            paragraph_format.first_line_indent = Inches(0.5)

                    p10 = document.add_paragraph()
                    run_p10 = p10.add_run(
                        '1.{} 工程建设批复情况：'.format(company_int + 1))
                    run_p10.bold = True

                    p11 = document.add_paragraph()
                    paragraph_format = p11.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    run_p11 = p11.add_run('1、批复情况：')
                    run_p11.bold = True
                    run_p11_2 = p11.add_run('本报告涉及立项批复内容较多，且均为会议纪要，在报告中不予详书。')

                    p12 = document.add_paragraph()
                    paragraph_format = p12.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    run_p12 = p12.add_run('2、工程建设单位：')
                    run_p12.bold = True
                    run_p12_2 = p12.add_run('中国移动通信集团内蒙古有限公司包头分公司')

                    p13 = document.add_paragraph()
                    paragraph_format = p13.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    run_p13 = p13.add_run('3、工程建设单位：')
                    run_p13.bold = True
                    run_p13_2 = p13.add_run('无')

                    document.add_heading('二 、审计内容', level=2)

                    p14 = document.add_paragraph()
                    paragraph_format = p14.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p14.add_run('1、结算方式是否符合合同约定。')

                    p15 = document.add_paragraph()
                    paragraph_format = p15.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p15.add_run('2、结算工程量与实际工程量的误差情况。')

                    p16 = document.add_paragraph()
                    paragraph_format = p16.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p16.add_run('3、工程量是否符合规定的计算规则，数量是否准确。')

                    p17 = document.add_paragraph()
                    paragraph_format = p17.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p17.add_run('4、结算单价的套用是否合理。')

                    p18 = document.add_paragraph()
                    paragraph_format = p18.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p18.add_run('5、工程取费是否执行相应的计算基数和费率标准等。')

                    document.add_heading('三、审计依据', level=2)

                    p19 = document.add_paragraph()
                    paragraph_format = p19.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p19.add_run(
                        '1、中国移动通信集团内蒙古有限公司与湖南宏达工程造价咨询有限公司签订的《中国移动通信集团公司内蒙古有限公司委托建设项目结算审计框架合同》合同'
                    )

                    p20 = document.add_paragraph()
                    paragraph_format = p20.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p20.add_run('2、国家、行业相关文件及规定。')

                    p21 = document.add_paragraph()
                    paragraph_format = p21.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p21.add_run('3、建设单位提供的资料：')

                    p22 = document.add_paragraph()
                    paragraph_format = p22.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p22.add_run(' （1）原报工程结算书；')

                    p23 = document.add_paragraph()
                    paragraph_format = p23.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p23.add_run(' （2）施工图；')

                    p24 = document.add_paragraph()
                    paragraph_format = p24.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p24.add_run(' （3）施工合同；')

                    p25 = document.add_paragraph()
                    paragraph_format = p25.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p25.add_run(' （4）竣工技术文件；')

                    p26 = document.add_paragraph()
                    paragraph_format = p26.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p26.add_run(' （5）工程签证资料；')

                    p27 = document.add_paragraph()
                    paragraph_format = p27.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p27.add_run(' （6）工程定价资料；')

                    p28 = document.add_paragraph()
                    paragraph_format = p28.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p28.add_run(' （7）投标书、招标文件、投标文件。')

                    p29 = document.add_paragraph()
                    paragraph_format = p29.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p29.add_run('4、审计定额依据：')

                    p30 = document.add_paragraph()
                    paragraph_format = p30.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p30.add_run(
                        '''《通信建设工程概预算编制配套法规文件汇编（2008.5）》、《通信建设工程价款结算办法》、《通信建设工程预算定额（全册）》、内蒙古包头市材料价格信息及移动公司关于工程结算的有关规定。'''
                    )

                    p31 = document.add_paragraph()
                    paragraph_format = p31.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p31.add_run('5、现场勘察记录。')

                    document.add_heading('四、双方责任', level=2)

                    p32 = document.add_paragraph()
                    run_32 = p32.add_run('4.1 建设单位责任')
                    run_32.bold = True

                    p33 = document.add_paragraph()
                    paragraph_format = p33.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p33.add_run(
                        '1、建设单位通过审计委托书向审计机构委托工程建设项目审计，并向审计机构提供工程审计的全部资料。')

                    p34 = document.add_paragraph()
                    paragraph_format = p34.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p34.add_run(
                        '2、建设单位应对开展审核工作给予充分的合作，提供必要的工作条件，并按审计机构要求，及时提供与委托内容相关的完整的建设项目有关文件、规定、合同、图纸、预结算等资料，并对其真实性、合法性和完整性负责。'
                    )

                    p35 = document.add_paragraph()
                    paragraph_format = p35.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p35.add_run(
                        '3、审核过程中需进一步核对原始资料或计算依据的，建设单位及其施工单位应提供必要的工作条件及合作，凡需施工单位签认的基础资料由建设单位协助办理。'
                    )

                    p33 = document.add_paragraph()
                    paragraph_format = p33.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p33.add_run(
                        '4、工程结算审核结果，凡需建设单位、审计机构双方和施工单位合议的，由建设单位组织协调，施工单位可以充分发表意见，对审核确认的事项，施工单位无理拒签、又不参加合议的，不影响审计机构出具审核报告，由此产生的后果由责任者承担。'
                    )

                    p34 = document.add_paragraph()
                    paragraph_format = p34.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p34.add_run(
                        '5、由于建设单位违反约定事项，所提供的审核资料不真实，不合法、不完整、不及时，致使审核结果不当或延期提交审核报告，由建设单位承担责任。'
                    )

                    p35 = document.add_paragraph()
                    paragraph_format = p35.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p35.add_run('6、建设单位应正确使用审计报告，由于使用不当所造成的后果，与审计机构无关。')

                    p36 = document.add_paragraph()
                    paragraph_format = p36.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p36.add_run('7、建设单位应按照约定的条件，及时足额支付审计费。')

                    p37 = document.add_paragraph()
                    run_37 = p37.add_run('4.2审计机构责任')
                    run_37.bold = True

                    p38 = document.add_paragraph()
                    paragraph_format = p38.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p38.add_run('1、工程结算审核范围')

                    p39 = document.add_paragraph()
                    paragraph_format = p39.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p39.add_run('2、施工方工程施工费、安装费等的审核。')

                    p40 = document.add_paragraph()
                    paragraph_format = p40.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p40.add_run(
                        '3、审计机构应按照法规和准则的要求，根据建设单位提供的各项建设项目文件和工程资料，严格实施主要的审核程序，发表审核意见，按约定的时间出具真实合法的审核报告。'
                    )

                    p41 = document.add_paragraph()
                    paragraph_format = p41.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p41.add_run(
                        '4、检查弊端不属于一般审计工作范围，但在审计过程如发现被审计单位在基建工程管理，财务管理和财产物质管理方面存在问题，导致有产生重大弊端的可能，审计机构应将其情况报告建设单位。在审计过程中，如发现被审计单位的内部控制有重大缺陷，应将情况报告建设单位。'
                    )

                    p42 = document.add_paragraph()
                    paragraph_format = p42.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p42.add_run(
                        '5、审计机构应按照约定的审计时限、安排足够的审计人员到达审计现场开展审计工作，如不能及时安排审计人员进行审计，则建设单位有权安排其他审计单位进行审计。'
                    )

                    p43 = document.add_paragraph()
                    paragraph_format = p43.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p43.add_run(
                        '6、现场审计：审计机构应对扩容设备安装、综合楼生产楼等土建（包括水、电、暖等）、机房装修及城域网扩容管道、零星工程等工程项目， 100%现场实测、对同一家施工单位施工的光缆线路等单价低、分布广的工程，现场勘查率必须达到20%以上；审计过程中如出现工程量方面的纠纷或施工单位提出要求，必须现场抽查。'
                    )

                    p44 = document.add_paragraph()
                    paragraph_format = p44.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p44.add_run(
                        '7、审计机构应按照中标要求出具审计报告，由于审核人员是采取事后审计方法，并受被审计单位内部控制固有的局限和其他客观因素的制约，以及审核人员在审核中可能未予发现的疏漏。因此，审计机构的审核难免存在某些重要的方面反映失实，但这不能替代、减轻或免除建设单位对工程结算的编制责任。'
                    )

                    p45 = document.add_paragraph()
                    paragraph_format = p45.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p45.add_run(
                        '8、审计机构应于每周一报送审计周报，对建设单位委托的审计项目的审计进度、存在问题等情况进行反馈。')

                    p46 = document.add_paragraph()
                    paragraph_format = p46.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p46.add_run(
                        '9、由于审计机构违反工程造价审核规定和程序，致使审核结果严重失真，审计机构应重新进行审核，由此产生的费用由审计机构自行负责，若给建设单位造成了经济损失，审计机构还应全额赔偿建设单位的损失。'
                    )

                    document.add_heading('五、审核过程', level=2)

                    p47 = document.add_paragraph()
                    paragraph_format = p47.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p47.add_run(
                        '我们受贵公司委托，于2017年3月1日至2018年1月15日，对中国移动通信集团内蒙古有限公司包头业务区{0}{1}等{2}项工程进行了审计。在审计过程中，我们采用全面审核的方法，结合工程的实际情况，实施了包括逐项审查结算资料、现场勘察和有关工程技术人员交换意见等我们认为必要的审核程序。'
                        .format(str_type, str_pact, type_int))

                    p48 = document.add_paragraph()
                    paragraph_format = p48.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p48.add_run('我们根据本项目的特点，制定了审计工作计划。按照审计计划，我们将审计工作分为三个阶段：')

                    p49 = document.add_paragraph()
                    paragraph_format = p49.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p49.add_run(
                        '第一阶段，熟悉工程情况、搜集审计资料工作，并在此期间进行了结算初步审计工作，形成审计明细初稿。')

                    p50 = document.add_paragraph()
                    paragraph_format = p50.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p50.add_run('搜集的主要资料有：')

                    p51 = document.add_paragraph()
                    paragraph_format = p51.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p51.add_run('①	工程批复文件、招投标文件；')

                    p52 = document.add_paragraph()
                    paragraph_format = p52.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p52.add_run('②	工程报审结算书；')

                    p53 = document.add_paragraph()
                    paragraph_format = p53.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p53.add_run('③	工程施工合同、设计文件；')

                    p54 = document.add_paragraph()
                    paragraph_format = p54.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p54.add_run('④	工程竣工技术资料（开工报告、完工报告、初验证书、各种涉造价部分文件等）；')

                    p55 = document.add_paragraph()
                    paragraph_format = p55.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p55.add_run(
                        '第二阶段，实施了100%现场勘察、结算核对工作，对现场安装工程量实行点单型核查，并与施工单位、监理单位、建设单位交换了意见，依据施工合同对结算方式进行了确认，达成共识后，确认工程结算审核定案表。'
                    )

                    p56 = document.add_paragraph()
                    paragraph_format = p56.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p56.add_run('第三阶段，形成审计资料终稿，拟订审计报告初稿，与各方交换意见后，出具正式审计报告。')

                    document.add_heading('六、审核结果', level=2)

                    p57 = document.add_paragraph()
                    run_57 = p57.add_run('6.1 经济性评价')
                    run_57.bold = True

                    p58 = document.add_paragraph()
                    paragraph_format = p58.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p58.add_run(
                        '6.1.1 本报告所含{0}项中国移动通信集团内蒙古有限公司包头业务区{1}等工程的具体情况如下（详见附件一：工程结算费用审核汇总表）：'
                        .format(type_int, str_type))

                    p59 = document.add_paragraph()
                    paragraph_format = p59.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p59.add_run(
                        '6.1.2 本报告{0}项工程分由{1}家施工单位实施，合同金额共为{2:.2f}元，总送审金额为{3:.2f}元，审计总造价为{4:.2f}元，审减造价{5:.2f}元，平均审减率为{6:.2f}%，审计费用共计{7}元。'
                        .format(type_int, len(company_list), zong_jing_e,
                                song_shen_e, shen_ding_e, sheng_jian,
                                avg_shen_jian, 10403.03))  # 这里有他妈大问题

                    p60 = document.add_paragraph()
                    paragraph_format = p60.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    run_60 = p60.add_run('6.1.3  这里有他妈大问题！')
                    run_60.font.size = Pt(40)

                    p61 = document.add_paragraph()
                    paragraph_format = p61.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    run_61 = p61.add_run('6.1.4  这里有他妈大问题！')

                    p62 = document.add_paragraph()
                    run_62 = p62.add_run('6.2 主要核减原因')
                    run_62.bold = True

                    p63 = document.add_paragraph()
                    paragraph_format = p63.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p63.add_run('1、送审工作量与现场工作量不符')

                    p64 = document.add_paragraph()
                    paragraph_format = p64.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p64.add_run(
                        '例如：商丘市国基建筑安装有限公司送审的青山赛音道渠道装修施工合同工程中，施工单位送审门头贴膜8.766m2，现场实际查勘是4.35m2，审计时都按现场测量情况予以审减。'
                    )

                    p65 = document.add_paragraph()
                    paragraph_format = p65.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    run_65 = p65.add_run('2、定额取值套用错误')

                    p66 = document.add_paragraph()
                    paragraph_format = p66.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p66.add_run(
                        '例如：山西胜唐通信建设有限公司送审的第二批社会渠道传输改造项目林荫路科技楼他建专营店等2处施工合同中，施工单位送审布放墙壁光缆套取定额子目TXL-049（架设吊线式墙壁光缆），实际现场查看施工单位布放光缆方式是利旧式墙壁吊线，应套取定额子目TXL3-183(城区架设架空光缆)。'
                    )

                    document.add_heading('七、审核问题', level=2)

                    p67 = document.add_paragraph()
                    paragraph_format = p67.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p67.add_run('本次审计共发现问题1个，提出1个管理建议。')

                    p68 = document.add_paragraph()
                    paragraph_format = p68.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p68.add_run(
                        '本报告所涉部分项目为成本类工程，全部采用定额计价法结算，施工单位如果没有专业预算人员，容易发生定额套用错误或者借用定额不符等争议，导致验收、竣工结算、审计时产生不便，既影响施工单位合法效益，又加大建设单位成本管理难度。'
                    )

                    p69 = document.add_paragraph()
                    paragraph_format = p69.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p69.add_run(
                        '管理建议：建议建设单位针对工艺不复杂，材料采购简单且价格容易调查的项目在年度招标时实施清单报价法计价，既方便双方清楚造价的构成又避免后期的争议。建设单位和监理只要负责施工单位施工使用的材料是否与招标要求相符，完成质量、工程量可以在验收是核定确认，方便、明了。'
                    )

                    document.add_heading('八、其他说明', level=2)

                    p70 = document.add_paragraph()
                    paragraph_format = p70.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p70.add_run(
                        '1.本报告中所涉建设单位提供的相关依据资料（批复立项资料、施工合同、开工报告、验收报告等）由于在电子档案中已存在，报告不予重复。'
                    )

                    p71 = document.add_paragraph()
                    paragraph_format = p71.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p71.add_run(
                        '2.本报告一式四份，中国移动通信集团内蒙古有限公司及包头分公司共持三份，湖南宏达工程造价咨询有限公司持一份。'
                    )

                    document.add_paragraph()
                    document.add_paragraph()
                    document.add_paragraph()
                    document.add_paragraph()
                    document.add_paragraph()
                    document.add_paragraph()
                    document.add_paragraph()
                    document.add_paragraph()

                    p72 = document.add_paragraph()
                    run_72 = p72.add_run('附件:')
                    run_72.bold = True

                    p73 = document.add_paragraph()
                    paragraph_format = p73.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p73.add_run('1、附件一：工程结算费用审核汇总表')

                    p74 = document.add_paragraph()
                    paragraph_format = p74.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p74.add_run('2、附件二：单项工程审计定案表')

                    p75 = document.add_paragraph()
                    paragraph_format = p75.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p75.add_run('3、附件三：单项工程审定明细表')

                    p76 = document.add_paragraph()
                    paragraph_format = p76.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p76.add_run('4、附件四：结算审计现场审核记录表')

                    p77 = document.add_paragraph()
                    paragraph_format = p77.paragraph_format
                    paragraph_format.first_line_indent = Inches(0.5)
                    p77.add_run('5、附件五：部分现场审计痕迹记录')

                    document.add_paragraph()
                    document.add_paragraph()
                    document.add_paragraph()

                    p78 = document.add_paragraph()
                    p78.add_run('湖南宏达工程造价咨询有限公司           项目负责人：付红卫')

                    p79 = document.add_paragraph()
                    p79.add_run(
                        '                                  编审： 彭创  李庆辉  梁湘')

                    p80 = document.add_paragraph()
                    p80.add_run(
                        '                                  复审： 陈晔  洪艳  胡红梅')

                    p81 = document.add_paragraph()
                    p81.add_run('地址：长沙市天心区湘府中路369号       审计公司执业印章')

                    p82 = document.add_paragraph()
                    p82.add_run('     星城荣域综合楼房1310')

                    p83 = document.add_paragraph()
                    p83.add_run('                                2018年1月25日')
                    name_int += 1
                    document.save('{0}\{1}{2}--{3}40.docx'.format(
                        en_doc.get(), str_type, name_int, last_item[-1][-1]))

                    list_box.insert(
                        0, '{0}:{1}-->40生成完成！'.format(str_type,
                                                      last_item[-1][-1]))

    lb_check_excel = Label(text="选择excel文件：", fg='blue')
    lb_check_excel.grid(row=0, column=1)

    lb_choose_dir = Label(text="选择存放路径：", fg='blue')
    lb_choose_dir.grid(row=1, column=1)

    en_excel = Entry(width=40, textvariable=excel_path)
    en_excel.grid(row=0, column=2)

    en_doc = Entry(width=40, textvariable=doc_path)
    en_doc.grid(row=1, column=2)

    bu_excel = Button(text="选择", width=5, command=get_excel)
    bu_excel.grid(row=0, column=3)

    bu_doc = Button(text='选择', width=5, command=tar_dir)
    bu_doc.grid(row=1, column=3)

    bu_check = Button(text='确定', width=5, command=to_docx)
    bu_check.grid(row=6, column=3)

    list_box = Listbox(width=60)
    list_box.grid(row=3, rowspan=3, columnspan=4)

    root.mainloop()


if __name__ == "__main__":
    handle_excel_gui()
